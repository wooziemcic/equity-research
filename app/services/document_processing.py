from __future__ import annotations

import csv
import hashlib
import io
import json
import re
import zipfile
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

from app import config
from app.services import processing_workspace


PARSER_VERSION = "5.0"
TEXT_EXTENSIONS = {".txt"}
HTML_EXTENSIONS = {".htm", ".html"}
CSV_EXTENSIONS = {".csv"}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xlsm"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}


@dataclass(frozen=True)
class ParsedPage:
    page_number: int
    page_label: str
    text: str
    extraction_method: str
    native_text_character_count: int = 0
    ocr_text_character_count: int = 0
    ocr_confidence: float | None = None
    page_text_path: str | None = None
    image_render_path: str | None = None
    warnings: list[str] = field(default_factory=list)


@dataclass(frozen=True)
class ParsedSheet:
    sheet_name: str
    sheet_index: int
    hidden_state: str
    used_range: str
    formula_cell_count: int
    cached_value_cell_count: int
    external_link_count: int
    warning_flags: list[str]
    extracted_representation_path: str | None = None


@dataclass(frozen=True)
class SourceSegment:
    text: str
    locator: dict[str, Any]
    extraction_method: str
    page_number: int | None = None
    sheet_name: str | None = None
    row_range: str | None = None
    section_heading: str | None = None


@dataclass(frozen=True)
class ParsedDocument:
    status: str
    parser_used: str
    parser_version: str
    detected_language: str
    pages: list[ParsedPage]
    sheets: list[ParsedSheet]
    segments: list[SourceSegment]
    full_text_path: str | None
    extracted_character_count: int
    page_count: int = 0
    sheet_count: int = 0
    table_count: int = 0
    ocr_required: bool = False
    ocr_pages: int = 0
    warnings: list[str] = field(default_factory=list)
    error_message: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ChunkDraft:
    chunk_text: str
    chunk_index: int
    character_count: int
    token_estimate: int
    extraction_method: str
    source_locator: dict[str, Any]
    chunk_hash: str
    page_number: int | None = None
    sheet_name: str | None = None
    row_range: str | None = None
    section_heading: str | None = None


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _safe_read_text(path: Path) -> tuple[str, str, list[str]]:
    warnings: list[str] = []
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding), encoding, warnings
        except UnicodeDecodeError:
            continue
    warnings.append("Text encoding could not be detected; decoded with latin-1 replacement.")
    return raw.decode("latin-1", errors="replace"), "latin-1-replace", warnings


def _write_full_text(workspace: Path, text: str) -> str:
    path = processing_workspace.atomic_write_text(workspace / "full_text.txt", text)
    return processing_workspace.relative_processed_path(path)


def _base_locator(version_doc: dict[str, Any], extraction_method: str) -> dict[str, Any]:
    return {
        "version_document_id": version_doc["document_id"],
        "original_document_id": version_doc.get("original_document_id"),
        "display_title": version_doc.get("title") or version_doc.get("package_filename"),
        "relative_package_path": version_doc["relative_package_path"],
        "extraction_method": extraction_method,
    }


def _trim_extracted_text(text: str, warnings: list[str]) -> str:
    if len(text) <= config.MAX_EXTRACTED_CHARACTERS:
        return text
    warnings.append("Extracted text exceeded MAX_EXTRACTED_CHARACTERS and was truncated.")
    return text[: config.MAX_EXTRACTED_CHARACTERS]


def parse_version_document(
    *,
    version_doc: dict[str, Any],
    source_path: Path,
    version_id: str,
    processing_run_id: str,
    ocr_enabled: bool | None = None,
) -> ParsedDocument:
    workspace = processing_workspace.document_workspace(version_id, processing_run_id, version_doc["document_id"])
    extension = source_path.suffix.lower()
    ocr = config.OCR_ENABLED if ocr_enabled is None else ocr_enabled
    try:
        if extension in TEXT_EXTENSIONS:
            parsed = _parse_text(version_doc, source_path, workspace)
        elif extension in HTML_EXTENSIONS:
            parsed = _parse_html(version_doc, source_path, workspace)
        elif extension in CSV_EXTENSIONS:
            parsed = _parse_csv(version_doc, source_path, workspace)
        elif extension in SPREADSHEET_EXTENSIONS:
            parsed = _parse_spreadsheet(version_doc, source_path, workspace)
        elif extension == ".docx":
            parsed = _parse_docx(version_doc, source_path, workspace)
        elif extension == ".pdf":
            parsed = _parse_pdf(version_doc, source_path, workspace, ocr_enabled=ocr)
        elif extension in IMAGE_EXTENSIONS:
            parsed = _parse_image(version_doc, source_path, workspace, ocr_enabled=ocr)
        elif extension == ".zip":
            parsed = _parse_zip(version_doc, source_path, workspace)
        else:
            parsed = _failed_parse(
                version_doc,
                workspace,
                parser_used="UNSUPPORTED",
                error=f"Unsupported Phase 5 parser for extension {extension}.",
            )
    except Exception as exc:
        parsed = _failed_parse(version_doc, workspace, parser_used=extension.lstrip(".").upper() or "UNKNOWN", error=str(exc))
    metadata_path = workspace / "document_metadata.json"
    processing_workspace.atomic_write_json(
        metadata_path,
        {
            "version_document_id": version_doc["document_id"],
            "relative_package_path": version_doc["relative_package_path"],
            "parser_used": parsed.parser_used,
            "parser_version": parsed.parser_version,
            "status": parsed.status,
            "warnings": parsed.warnings,
            "metadata": parsed.metadata,
        },
    )
    processing_workspace.atomic_write_json(workspace / "warnings.json", parsed.warnings)
    return parsed


def _failed_parse(
    version_doc: dict[str, Any],
    workspace: Path,
    *,
    parser_used: str,
    error: str,
) -> ParsedDocument:
    full_text_path = _write_full_text(workspace, "")
    return ParsedDocument(
        status=config.DOCUMENT_PROCESSING_FAILED,
        parser_used=parser_used,
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=[],
        sheets=[],
        segments=[],
        full_text_path=full_text_path,
        extracted_character_count=0,
        warnings=[],
        error_message=error,
        metadata={"document_id": version_doc["document_id"]},
    )


def _parse_text(version_doc: dict[str, Any], source_path: Path, workspace: Path) -> ParsedDocument:
    warnings: list[str] = []
    text, encoding, decode_warnings = _safe_read_text(source_path)
    warnings.extend(decode_warnings)
    text = _trim_extracted_text(text, warnings)
    lines = text.splitlines()
    page_path = processing_workspace.atomic_write_text(workspace / "pages" / "lines.txt", text)
    locator = _base_locator(version_doc, "NATIVE_TEXT")
    locator.update({"line_range": f"1-{len(lines)}", "source_text_hash": sha256_text(text)})
    segment = SourceSegment(
        text=text,
        locator=locator,
        extraction_method="NATIVE_TEXT",
        section_heading=_detect_section_heading(text),
    )
    page = ParsedPage(
        page_number=1,
        page_label="lines 1-%s" % len(lines),
        text=text,
        extraction_method="NATIVE_TEXT",
        native_text_character_count=len(text),
        page_text_path=processing_workspace.relative_processed_path(page_path),
        warnings=warnings,
    )
    return ParsedDocument(
        status=config.DOCUMENT_PROCESSING_SUCCESS,
        parser_used="TXT",
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=[page],
        sheets=[],
        segments=[segment] if text.strip() else [],
        full_text_path=_write_full_text(workspace, text),
        extracted_character_count=len(text),
        page_count=1,
        warnings=warnings,
        metadata={"encoding": encoding, "line_count": len(lines)},
    )


class _VisibleHTMLTextParser(HTMLParser):
    _SKIPPED = {"script", "style", "noscript", "template"}
    _BLOCKS = {"br", "p", "div", "tr", "li", "table", "section", "article", "h1", "h2", "h3", "h4", "h5", "h6"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        del attrs
        lowered = tag.lower()
        if lowered in self._SKIPPED:
            self.skip_depth += 1
        elif lowered in self._BLOCKS and not self.skip_depth:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        lowered = tag.lower()
        if lowered in self._SKIPPED and self.skip_depth:
            self.skip_depth -= 1
        elif lowered in self._BLOCKS and not self.skip_depth:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.skip_depth and data.strip():
            self.parts.append(data)

    def text(self) -> str:
        return "\n".join(line for line in (normalize_text(part) for part in self.parts) if line)


def _parse_html(version_doc: dict[str, Any], source_path: Path, workspace: Path) -> ParsedDocument:
    warnings: list[str] = []
    source, encoding, decode_warnings = _safe_read_text(source_path)
    warnings.extend(decode_warnings)
    parser = _VisibleHTMLTextParser()
    parser.feed(source)
    text = _trim_extracted_text(parser.text(), warnings)
    lines = text.splitlines()
    page_path = processing_workspace.atomic_write_text(workspace / "pages" / "html_text.txt", text)
    locator = _base_locator(version_doc, "HTML_TEXT")
    locator.update({"line_range": f"1-{len(lines)}", "source_text_hash": sha256_text(text)})
    segment = SourceSegment(
        text=text,
        locator=locator,
        extraction_method="HTML_TEXT",
        section_heading=_detect_section_heading(text),
    )
    page = ParsedPage(
        page_number=1,
        page_label=f"HTML lines 1-{len(lines)}",
        text=text,
        extraction_method="HTML_TEXT",
        native_text_character_count=len(text),
        page_text_path=processing_workspace.relative_processed_path(page_path),
        warnings=warnings,
    )
    return ParsedDocument(
        status=config.DOCUMENT_PROCESSING_SUCCESS if text.strip() else config.DOCUMENT_PROCESSING_PARTIAL,
        parser_used="HTML",
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=[page] if text.strip() else [],
        sheets=[],
        segments=[segment] if text.strip() else [],
        full_text_path=_write_full_text(workspace, text),
        extracted_character_count=len(text),
        page_count=1 if text.strip() else 0,
        warnings=warnings + ([] if text.strip() else ["HTML document contained no visible text."]),
        metadata={"encoding": encoding, "line_count": len(lines), "source_format": "HTML"},
    )


def _parse_csv(version_doc: dict[str, Any], source_path: Path, workspace: Path) -> ParsedDocument:
    warnings: list[str] = []
    text, encoding, decode_warnings = _safe_read_text(source_path)
    warnings.extend(decode_warnings)
    sample = text[:2048]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel
        warnings.append("CSV delimiter could not be detected; comma delimiter assumed.")
    rows = list(csv.reader(io.StringIO(text), dialect))
    headers = rows[0] if rows else []
    segments: list[SourceSegment] = []
    rendered_rows: list[str] = []
    for row_index, row in enumerate(rows, start=1):
        if row_index == 1 and headers:
            rendered = "Headers: " + " | ".join(headers)
        elif headers:
            cells = [f"{headers[i] if i < len(headers) else get_column_letter(i + 1)}={value}" for i, value in enumerate(row)]
            rendered = f"Row {row_index}: " + " | ".join(cells)
        else:
            rendered = f"Row {row_index}: " + " | ".join(row)
        rendered_rows.append(rendered)
        locator = _base_locator(version_doc, "CSV")
        locator.update({"row_range": str(row_index), "source_text_hash": sha256_text(rendered)})
        segments.append(SourceSegment(text=rendered, locator=locator, extraction_method="CSV", row_range=str(row_index)))
    full_text = _trim_extracted_text("\n".join(rendered_rows), warnings)
    rows_path = processing_workspace.atomic_write_text(workspace / "pages" / "rows.txt", full_text)
    page = ParsedPage(
        page_number=1,
        page_label=f"rows 1-{len(rows)}",
        text=full_text,
        extraction_method="CSV",
        native_text_character_count=len(full_text),
        page_text_path=processing_workspace.relative_processed_path(rows_path),
        warnings=warnings,
    )
    return ParsedDocument(
        status=config.DOCUMENT_PROCESSING_SUCCESS,
        parser_used="CSV",
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=[page],
        sheets=[],
        segments=segments,
        full_text_path=_write_full_text(workspace, full_text),
        extracted_character_count=len(full_text),
        page_count=1,
        table_count=1 if rows else 0,
        warnings=warnings,
        metadata={"encoding": encoding, "delimiter": getattr(dialect, "delimiter", ","), "row_count": len(rows), "headers": headers},
    )


def _parse_spreadsheet(version_doc: dict[str, Any], source_path: Path, workspace: Path) -> ParsedDocument:
    warnings: list[str] = []
    workbook = load_workbook(source_path, data_only=False, read_only=False, keep_vba=False, keep_links=False)
    values_workbook = load_workbook(source_path, data_only=True, read_only=True, keep_vba=False, keep_links=False)
    external_link_count = len(getattr(workbook, "_external_links", []) or [])
    if source_path.suffix.lower() == ".xlsm":
        warnings.append("XLSM opened without executing macros.")
    if external_link_count:
        warnings.append("Workbook contains external link metadata; links were not followed.")
    sheet_names = workbook.sheetnames[: config.MAX_SPREADSHEET_SHEETS]
    if len(workbook.sheetnames) > len(sheet_names):
        warnings.append("Workbook exceeded MAX_SPREADSHEET_SHEETS; remaining sheets were skipped.")
    sheets: list[ParsedSheet] = []
    segments: list[SourceSegment] = []
    rendered_parts: list[str] = []
    cells_seen = 0
    for sheet_index, sheet_name in enumerate(sheet_names, start=1):
        ws = workbook[sheet_name]
        values_ws = values_workbook[sheet_name]
        max_row = max(ws.max_row or 1, 1)
        max_column = max(ws.max_column or 1, 1)
        used_range = f"A1:{get_column_letter(max_column)}{max_row}"
        formula_count = 0
        cached_formula_count = 0
        warning_flags: list[str] = []
        sheet_lines: list[str] = []
        row_start = 1
        for row_index, row in enumerate(ws.iter_rows(), start=1):
            row_parts: list[str] = []
            for cell in row:
                cells_seen += 1
                if cells_seen > config.MAX_SPREADSHEET_CELLS:
                    warning_flags.append("MAX_SPREADSHEET_CELLS_REACHED")
                    warnings.append("Workbook exceeded MAX_SPREADSHEET_CELLS; remaining cells were skipped.")
                    break
                value = cell.value
                label = "UNKNOWN"
                rendered_value = "" if value is None else str(value)
                if isinstance(value, str) and value.startswith("="):
                    formula_count += 1
                    cached_value = values_ws[cell.coordinate].value
                    if cached_value is None:
                        label = "FORMULA_WITHOUT_CACHED_VALUE"
                        warning_flags.append("FORMULA_WITHOUT_CACHED_VALUE")
                    else:
                        label = "CACHED_FORMULA_VALUE"
                        cached_formula_count += 1
                        rendered_value = f"{value} cached={cached_value}"
                        warning_flags.append("CACHED_FORMULA_VALUE")
                elif value is not None:
                    label = "LITERAL_VALUE"
                if value is not None:
                    row_parts.append(f"{cell.coordinate}={rendered_value} [{label}]")
            if row_parts:
                sheet_lines.append(f"Row {row_index}: " + " | ".join(row_parts))
            if "MAX_SPREADSHEET_CELLS_REACHED" in warning_flags:
                break
        if "CACHED_FORMULA_VALUE" in warning_flags:
            warnings.append(f"Sheet {sheet_name} contains cached formula values that may be stale.")
        if "FORMULA_WITHOUT_CACHED_VALUE" in warning_flags:
            warnings.append(f"Sheet {sheet_name} contains formulas without cached values.")
        sheet_text = "\n".join(sheet_lines)
        path = processing_workspace.atomic_write_text(workspace / "sheets" / f"{_safe_name(sheet_name)}.txt", sheet_text)
        sheets.append(
            ParsedSheet(
                sheet_name=sheet_name,
                sheet_index=sheet_index,
                hidden_state=ws.sheet_state,
                used_range=used_range,
                formula_cell_count=formula_count,
                cached_value_cell_count=cached_formula_count,
                external_link_count=external_link_count,
                warning_flags=sorted(set(warning_flags)),
                extracted_representation_path=processing_workspace.relative_processed_path(path),
            )
        )
        rendered_parts.append(f"Sheet {sheet_name} ({used_range})\n{sheet_text}")
        if sheet_text.strip():
            row_range = f"{row_start}-{max_row}"
            locator = _base_locator(version_doc, "SPREADSHEET")
            locator.update(
                {
                    "sheet_name": sheet_name,
                    "cell_range": used_range,
                    "row_range": row_range,
                    "source_text_hash": sha256_text(sheet_text),
                }
            )
            segments.append(
                SourceSegment(
                    text=sheet_text,
                    locator=locator,
                    extraction_method="SPREADSHEET",
                    sheet_name=sheet_name,
                    row_range=row_range,
                    section_heading=sheet_name,
                )
            )
        if cells_seen > config.MAX_SPREADSHEET_CELLS:
            break
    full_text = _trim_extracted_text("\n\n".join(rendered_parts), warnings)
    status = config.DOCUMENT_PROCESSING_PARTIAL if any("MAX_" in warning for warning in warnings) else config.DOCUMENT_PROCESSING_SUCCESS
    workbook.close()
    values_workbook.close()
    return ParsedDocument(
        status=status,
        parser_used=source_path.suffix.lower().lstrip(".").upper(),
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=[],
        sheets=sheets,
        segments=segments,
        full_text_path=_write_full_text(workspace, full_text),
        extracted_character_count=len(full_text),
        sheet_count=len(sheets),
        table_count=len([sheet for sheet in sheets if sheet.used_range != "A1:A1"]),
        warnings=sorted(set(warnings)),
        metadata={"sheet_names": workbook.sheetnames, "external_link_count": external_link_count, "macros_executed": False},
    )


def _parse_docx(version_doc: dict[str, Any], source_path: Path, workspace: Path) -> ParsedDocument:
    import docx

    warnings: list[str] = []
    document = docx.Document(source_path)
    segments: list[SourceSegment] = []
    lines: list[str] = []
    current_heading: str | None = None
    paragraph_index = 0
    for paragraph in document.paragraphs:
        paragraph_index += 1
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            current_heading = text
            lines.append(text)
            continue
        rendered = f"Paragraph {paragraph_index}: {text}"
        lines.append(rendered)
        locator = _base_locator(version_doc, "DOCX")
        locator.update({"section_heading": current_heading, "paragraph": paragraph_index, "source_text_hash": sha256_text(text)})
        segments.append(SourceSegment(text=text, locator=locator, extraction_method="DOCX", section_heading=current_heading))
    table_count = 0
    for table_index, table in enumerate(document.tables, start=1):
        table_count += 1
        table_rows: list[str] = []
        for row_index, row in enumerate(table.rows, start=1):
            row_text = " | ".join(normalize_text(cell.text) for cell in row.cells)
            if row_text.strip():
                table_rows.append(f"Table {table_index} Row {row_index}: {row_text}")
        table_text = "\n".join(table_rows)
        if table_text.strip():
            lines.append(table_text)
            locator = _base_locator(version_doc, "DOCX_TABLE")
            locator.update({"section_heading": current_heading, "table_index": table_index, "row_range": f"1-{len(table_rows)}", "source_text_hash": sha256_text(table_text)})
            segments.append(SourceSegment(text=table_text, locator=locator, extraction_method="DOCX_TABLE", row_range=f"1-{len(table_rows)}", section_heading=current_heading))
    full_text = _trim_extracted_text("\n".join(lines), warnings)
    page_path = processing_workspace.atomic_write_text(workspace / "pages" / "docx_sections.txt", full_text)
    page = ParsedPage(
        page_number=1,
        page_label="DOCX sections",
        text=full_text,
        extraction_method="DOCX",
        native_text_character_count=len(full_text),
        page_text_path=processing_workspace.relative_processed_path(page_path),
        warnings=warnings,
    )
    return ParsedDocument(
        status=config.DOCUMENT_PROCESSING_SUCCESS,
        parser_used="DOCX",
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=[page],
        sheets=[],
        segments=segments,
        full_text_path=_write_full_text(workspace, full_text),
        extracted_character_count=len(full_text),
        page_count=1,
        table_count=table_count,
        warnings=warnings,
        metadata={"paragraph_count": paragraph_index, "table_count": table_count, "page_numbers_reliable": False},
    )


def _parse_pdf(version_doc: dict[str, Any], source_path: Path, workspace: Path, *, ocr_enabled: bool) -> ParsedDocument:
    warnings: list[str] = []
    try:
        import fitz
    except ImportError:
        return _parse_pdf_with_pypdf(version_doc, source_path, workspace)
    pages: list[ParsedPage] = []
    segments: list[SourceSegment] = []
    ocr_required = False
    ocr_pages = 0
    page_texts: list[str] = []
    with fitz.open(source_path) as pdf:
        metadata = dict(pdf.metadata or {})
        page_total = pdf.page_count
        if page_total > config.MAX_PDF_PAGES:
            warnings.append("PDF exceeded MAX_PDF_PAGES; remaining pages were skipped.")
        for page_index in range(min(page_total, config.MAX_PDF_PAGES)):
            page = pdf.load_page(page_index)
            native_text = page.get_text("text") or ""
            native_text = native_text.strip()
            image_count = len(page.get_images(full=True))
            method = "NATIVE_TEXT" if native_text else "IMAGE_ONLY"
            page_warnings: list[str] = []
            ocr_text = ""
            ocr_confidence: float | None = None
            if native_text and image_count:
                method = "MIXED"
            if len(native_text) < 20:
                ocr_required = True
                if ocr_enabled and ocr_pages < config.MAX_OCR_PAGES:
                    ocr_text, ocr_confidence = _ocr_pdf_page(page, workspace, page_index + 1, page_warnings)
                    if ocr_text:
                        method = "OCR" if not native_text else "MIXED_OCR"
                        ocr_pages += 1
                    else:
                        page_warnings.append("OCR requested but no text was produced.")
                elif ocr_enabled:
                    page_warnings.append("OCR page limit reached; page skipped.")
                else:
                    page_warnings.append("OCR required but disabled.")
            page_text = "\n".join(part for part in (native_text, ocr_text) if part).strip()
            page_texts.append(page_text)
            page_path = processing_workspace.atomic_write_text(workspace / "pages" / f"page_{page_index + 1:04d}.txt", page_text)
            page_label = str(page_index + 1)
            parsed_page = ParsedPage(
                page_number=page_index + 1,
                page_label=page_label,
                text=page_text,
                extraction_method=method,
                native_text_character_count=len(native_text),
                ocr_text_character_count=len(ocr_text),
                ocr_confidence=ocr_confidence,
                page_text_path=processing_workspace.relative_processed_path(page_path),
                warnings=page_warnings,
            )
            pages.append(parsed_page)
            warnings.extend(page_warnings)
            if page_text.strip():
                locator = _base_locator(version_doc, method)
                locator.update({"page_number": page_index + 1, "page_label": page_label, "source_text_hash": sha256_text(page_text)})
                segments.append(
                    SourceSegment(
                        text=page_text,
                        locator=locator,
                        extraction_method=method,
                        page_number=page_index + 1,
                        section_heading=_detect_section_heading(page_text),
                    )
                )
    full_text = _trim_extracted_text("\n\n".join(page_texts), warnings)
    if not pages:
        status = config.DOCUMENT_PROCESSING_PARTIAL
        warnings.append("PDF contained no pages.")
    elif any(page.extraction_method in {"IMAGE_ONLY"} for page in pages) or ocr_required:
        status = config.DOCUMENT_PROCESSING_PARTIAL if not full_text.strip() or warnings else config.DOCUMENT_PROCESSING_SUCCESS
    else:
        status = config.DOCUMENT_PROCESSING_SUCCESS
    parser_used = "PDF_NATIVE"
    if pages and all(page.extraction_method == "IMAGE_ONLY" for page in pages):
        parser_used = "PDF_IMAGE_ONLY"
    elif pages and any(page.extraction_method in {"MIXED", "MIXED_OCR"} for page in pages):
        parser_used = "PDF_MIXED"
    return ParsedDocument(
        status=status,
        parser_used=parser_used,
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=pages,
        sheets=[],
        segments=segments,
        full_text_path=_write_full_text(workspace, full_text),
        extracted_character_count=len(full_text),
        page_count=len(pages),
        ocr_required=ocr_required,
        ocr_pages=ocr_pages,
        warnings=warnings,
        metadata=metadata,
    )


def _parse_pdf_with_pypdf(version_doc: dict[str, Any], source_path: Path, workspace: Path) -> ParsedDocument:
    from pypdf import PdfReader

    warnings: list[str] = ["PyMuPDF unavailable; pypdf fallback used."]
    reader = PdfReader(str(source_path))
    pages: list[ParsedPage] = []
    segments: list[SourceSegment] = []
    page_texts: list[str] = []
    for index, page in enumerate(reader.pages[: config.MAX_PDF_PAGES], start=1):
        text = page.extract_text() or ""
        page_texts.append(text)
        page_path = processing_workspace.atomic_write_text(workspace / "pages" / f"page_{index:04d}.txt", text)
        pages.append(
            ParsedPage(
                page_number=index,
                page_label=str(index),
                text=text,
                extraction_method="NATIVE_TEXT" if text.strip() else "IMAGE_ONLY",
                native_text_character_count=len(text),
                page_text_path=processing_workspace.relative_processed_path(page_path),
                warnings=[] if text.strip() else ["OCR required but disabled."],
            )
        )
        if text.strip():
            locator = _base_locator(version_doc, "NATIVE_TEXT")
            locator.update({"page_number": index, "page_label": str(index), "source_text_hash": sha256_text(text)})
            segments.append(SourceSegment(text=text, locator=locator, extraction_method="NATIVE_TEXT", page_number=index))
    full_text = _trim_extracted_text("\n\n".join(page_texts), warnings)
    ocr_required = any(not page.text.strip() for page in pages)
    return ParsedDocument(
        status=config.DOCUMENT_PROCESSING_PARTIAL if ocr_required else config.DOCUMENT_PROCESSING_SUCCESS,
        parser_used="PDF_PYPDF",
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=pages,
        sheets=[],
        segments=segments,
        full_text_path=_write_full_text(workspace, full_text),
        extracted_character_count=len(full_text),
        page_count=len(pages),
        ocr_required=ocr_required,
        warnings=warnings,
        metadata={"metadata": getattr(reader, "metadata", {})},
    )


def _ocr_pdf_page(page: Any, workspace: Path, page_number: int, warnings: list[str]) -> tuple[str, float | None]:
    try:
        import pytesseract
        from PIL import Image

        pixmap = page.get_pixmap(matrix=None, alpha=False)
        image = Image.open(io.BytesIO(pixmap.tobytes("png")))
        image_path = processing_workspace.atomic_write_bytes(workspace / "pages" / f"page_{page_number:04d}.png", pixmap.tobytes("png"))
        warnings.append(f"OCR image rendered to {processing_workspace.relative_processed_path(image_path)}.")
        data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        words = [word for word in data.get("text", []) if str(word).strip()]
        confidences = [
            float(conf)
            for conf in data.get("conf", [])
            if str(conf).strip() not in {"", "-1"}
        ]
        confidence = (sum(confidences) / len(confidences) / 100.0) if confidences else None
        if confidence is not None and confidence < config.OCR_CONFIDENCE_THRESHOLD:
            warnings.append("OCR confidence below threshold; extracted numbers need review.")
        return " ".join(words), confidence
    except Exception as exc:
        warnings.append(f"OCR dependencies unavailable or failed locally: {exc}")
        return "", None


def _parse_image(version_doc: dict[str, Any], source_path: Path, workspace: Path, *, ocr_enabled: bool) -> ParsedDocument:
    from PIL import Image

    warnings: list[str] = []
    with Image.open(source_path) as image:
        metadata = {"width": image.width, "height": image.height, "mode": image.mode}
        text = ""
        confidence: float | None = None
        ocr_required = True
        if ocr_enabled:
            try:
                import pytesseract

                data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                words = [word for word in data.get("text", []) if str(word).strip()]
                text = " ".join(words)
                confidences = [
                    float(conf)
                    for conf in data.get("conf", [])
                    if str(conf).strip() not in {"", "-1"}
                ]
                confidence = (sum(confidences) / len(confidences) / 100.0) if confidences else None
            except Exception as exc:
                warnings.append(f"OCR dependencies unavailable or failed locally: {exc}")
        else:
            warnings.append("OCR required but disabled.")
    page_path = processing_workspace.atomic_write_text(workspace / "pages" / "image_ocr.txt", text)
    page = ParsedPage(
        page_number=1,
        page_label="image",
        text=text,
        extraction_method="OCR" if text else "IMAGE_ONLY",
        native_text_character_count=0,
        ocr_text_character_count=len(text),
        ocr_confidence=confidence,
        page_text_path=processing_workspace.relative_processed_path(page_path),
        warnings=warnings,
    )
    segments: list[SourceSegment] = []
    if text.strip():
        locator = _base_locator(version_doc, "OCR")
        locator.update({"page_number": 1, "source_text_hash": sha256_text(text)})
        segments.append(SourceSegment(text=text, locator=locator, extraction_method="OCR", page_number=1))
    return ParsedDocument(
        status=config.DOCUMENT_PROCESSING_PARTIAL if ocr_required and not text else config.DOCUMENT_PROCESSING_SUCCESS,
        parser_used="IMAGE",
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=[page],
        sheets=[],
        segments=segments,
        full_text_path=_write_full_text(workspace, text),
        extracted_character_count=len(text),
        page_count=1,
        ocr_required=ocr_required,
        ocr_pages=1 if text else 0,
        warnings=warnings,
        metadata=metadata,
    )


def _parse_zip(version_doc: dict[str, Any], source_path: Path, workspace: Path) -> ParsedDocument:
    warnings = ["NOT_PROCESSABLE_ARCHIVE: ZIP contents are not processed automatically in Phase 5."]
    metadata: dict[str, Any] = {"archive_only": True, "entries": []}
    try:
        with zipfile.ZipFile(source_path) as archive:
            metadata["entries"] = [
                {"filename": info.filename, "file_size": info.file_size, "compress_size": info.compress_size}
                for info in archive.infolist()
            ]
    except zipfile.BadZipFile:
        warnings.append("ZIP file is malformed.")
    processing_workspace.atomic_write_json(workspace / "document_metadata.json", metadata)
    return ParsedDocument(
        status=config.DOCUMENT_PROCESSING_SKIPPED,
        parser_used="ZIP_ARCHIVE_ONLY",
        parser_version=PARSER_VERSION,
        detected_language="unknown",
        pages=[],
        sheets=[],
        segments=[],
        full_text_path=_write_full_text(workspace, ""),
        extracted_character_count=0,
        warnings=warnings,
        metadata=metadata,
    )


def _safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._") or "sheet"


def _detect_section_heading(text: str) -> str | None:
    for line in text.splitlines():
        clean = normalize_text(line)
        if re.match(r"(?i)^item\s+\d+[a-z]?", clean):
            return clean[:120]
        if clean and clean == clean.upper() and 6 <= len(clean) <= 120:
            return clean
    return None


def chunk_parsed_document(parsed: ParsedDocument, *, chunk_size: int | None = None, overlap: int | None = None) -> list[ChunkDraft]:
    size = chunk_size or config.CHUNK_SIZE
    step_back = min(overlap if overlap is not None else config.CHUNK_OVERLAP, max(size - 1, 0))
    chunks: list[ChunkDraft] = []
    chunk_index = 0
    for segment in parsed.segments:
        text = normalize_text(segment.text)
        if not text:
            continue
        start = 0
        while start < len(text):
            end = min(start + size, len(text))
            if end < len(text):
                split_at = text.rfind(" ", start, end)
                if split_at > start + max(size // 2, 1):
                    end = split_at
            chunk_text = text[start:end].strip()
            if not chunk_text:
                break
            locator = dict(segment.locator)
            locator.update(
                {
                    "char_start": start,
                    "char_end": end,
                    "chunk_source_text_hash": sha256_text(chunk_text),
                }
            )
            chunk_index += 1
            chunks.append(
                ChunkDraft(
                    chunk_text=chunk_text,
                    chunk_index=chunk_index,
                    character_count=len(chunk_text),
                    token_estimate=max(1, len(chunk_text.split())),
                    extraction_method=segment.extraction_method,
                    source_locator=locator,
                    chunk_hash=sha256_text(normalize_text(chunk_text).lower()),
                    page_number=segment.page_number,
                    sheet_name=segment.sheet_name,
                    row_range=segment.row_range,
                    section_heading=segment.section_heading,
                )
            )
            if end >= len(text):
                break
            start = max(end - step_back, start + 1)
    return chunks
