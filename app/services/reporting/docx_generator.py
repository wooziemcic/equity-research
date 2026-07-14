from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def build_docx_report(path: Path, sections: list[dict[str, Any]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    header = section.header.paragraphs[0]
    header.text = "CUTLER EQUITY RESEARCH"
    footer = section.footer.paragraphs[0]
    footer.text = "Evidence-grounded investment memo for analyst and portfolio-manager review."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.space_after = Pt(3)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 1"].font.color.rgb = None
    styles["Heading 1"].paragraph_format.space_before = Pt(5)
    styles["Heading 1"].paragraph_format.space_after = Pt(2)
    for section_payload in sections:
        title = section_payload.get("title")
        if title:
            document.add_heading(str(title), level=1)
        for paragraph in section_payload.get("paragraphs", []):
            document.add_paragraph(str(paragraph))
        for table_payload in section_payload.get("tables", []):
            rows = table_payload.get("rows", [])
            if not rows:
                continue
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for col_index, value in enumerate(row):
                    cell = table.cell(row_index, col_index)
                    cell.text = "" if value is None else str(value)
                    if row_index == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
        if section_payload.get("page_break"):
            document.add_page_break()
    document.save(path)
