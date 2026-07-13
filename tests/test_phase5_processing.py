from __future__ import annotations

import io
import json
import zipfile
from datetime import date
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw
from streamlit.testing.v1 import AppTest

from app import config
from app.services.checklist_service import ensure_package_checklist
from app.services.document_processing import parse_version_document
from app.services.evidence_service import (
    create_analyst_evidence_from_chunk,
    detect_claim_conflicts,
    verify_evidence_record,
)
from app.services.package_builder import build_package_version, lock_version, sha256_file
from app.services.package_service import PackageInput, create_package
from app.services.processing_pipeline import run_processing_pipeline, validate_processing_eligibility
from app.services.retrieval_service import search_chunks
from app.services.upload_service import UploadCandidate, store_uploaded_files
from app.utils import database


@pytest.fixture(autouse=True)
def phase5_config(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(config, "DOWNLOAD_DIR", tmp_path / "downloaded")
    monkeypatch.setattr(config, "PACKAGE_DIR", tmp_path / "packages")
    monkeypatch.setattr(config, "PROCESSED_DIR", tmp_path / "processed")
    monkeypatch.setattr(config, "MAX_UPLOAD_FILE_MB", 20)
    monkeypatch.setattr(config, "MAX_UPLOAD_BATCH_MB", 100)
    monkeypatch.setattr(config, "MAX_PDF_PAGES", 20)
    monkeypatch.setattr(config, "MAX_SPREADSHEET_SHEETS", 10)
    monkeypatch.setattr(config, "MAX_SPREADSHEET_CELLS", 1000)
    monkeypatch.setattr(config, "MAX_EXTRACTED_CHARACTERS", 250000)
    monkeypatch.setattr(config, "OCR_ENABLED", False)
    monkeypatch.setattr(config, "MAX_OCR_PAGES", 2)
    monkeypatch.setattr(config, "CHUNK_SIZE", 600)
    monkeypatch.setattr(config, "CHUNK_OVERLAP", 60)


@pytest.fixture()
def temp_db(tmp_path: Path) -> Path:
    db_path = tmp_path / "phase5.db"
    database.initialize_database(db_path)
    return db_path


def _package(temp_db: Path, ticker: str = "QXO") -> dict:
    package = create_package(
        PackageInput(ticker, "Common Equity", date(2026, 7, 13), 3, ""),
        db_path=temp_db,
    )
    package = database.update_package_company_metadata(
        package["package_id"],
        {
            "ticker": ticker,
            "company_name": f"{ticker} Inc.",
            "cik": "0001234567",
            "exchange": "NYSE",
            "sic": "7370",
            "industry_description": "Services",
            "fiscal_year_end": "1231",
            "sec_company_url": "https://www.sec.gov/edgar/browse/?CIK=0001234567",
            "resolution_status": "RESOLVED",
            "resolution_source": "test",
            "resolution_timestamp": "2026-07-13T00:00:00+00:00",
        },
        db_path=temp_db,
    )
    ensure_package_checklist(package, db_path=temp_db)
    return database.update_package_review_acknowledgement(
        package["package_id"],
        checklist_reviewed=True,
        reviewed_by="analyst",
        review_note="Reviewed for Phase 5 tests.",
        missing_core_acknowledged=True,
        stale_documents_acknowledged=True,
        needs_review_acknowledged=True,
        db_path=temp_db,
    )


def _pdf_bytes(pages: list[str]) -> bytes:
    import fitz

    pdf = fitz.open()
    for text in pages:
        page = pdf.new_page()
        page.insert_text((72, 72), text, fontsize=11)
    return pdf.tobytes()


def _image_only_pdf_bytes() -> bytes:
    import fitz

    image = Image.new("RGB", (300, 120), "white")
    draw = ImageDraw.Draw(image)
    draw.text((10, 45), "Image only revenue $555 million", fill="black")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    pdf = fitz.open()
    page = pdf.new_page(width=300, height=120)
    page.insert_image(page.rect, stream=buffer.getvalue())
    return pdf.tobytes()


def _mixed_pdf_bytes() -> bytes:
    import fitz

    image = Image.new("RGB", (120, 60), "white")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    pdf = fitz.open()
    page = pdf.new_page(width=300, height=140)
    page.insert_text((72, 72), "Mixed PDF revenue was $321 million in FY2026.", fontsize=11)
    page.insert_image(fitz.Rect(20, 80, 140, 130), stream=buffer.getvalue())
    return pdf.tobytes()


def _docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Liquidity", level=1)
    doc.add_paragraph("Liquidity was $200 million in FY2026.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Debt"
    table.cell(1, 1).text = "$50 million"
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _xlsx_bytes(*, hidden: bool = True) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Estimates"
    sheet["A1"] = "Metric"
    sheet["B1"] = "FY2026"
    sheet["A2"] = "Revenue"
    sheet["B2"] = 123
    sheet["C2"] = "=B2*2"
    sheet["A3"] = "Adjusted margin"
    sheet["B3"] = "15%"
    if hidden:
        hidden_sheet = workbook.create_sheet("HiddenData")
        hidden_sheet.sheet_state = "hidden"
        hidden_sheet["A1"] = "Debt"
        hidden_sheet["B1"] = 99
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _zip_bytes() -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("inside.txt", "Revenue was $1 million.")
    return buffer.getvalue()


def _lock_version_with_uploads(temp_db: Path, uploads: list[UploadCandidate], metadata: dict[str, dict] | None = None, ticker: str = "QXO") -> dict:
    package = _package(temp_db, ticker=ticker)
    store_uploaded_files(
        package,
        uploads,
        source_type="other",
        authorization_confirmed=True,
        metadata_by_name=metadata or {candidate.original_filename: {"final_category_code": "other"} for candidate in uploads},
        db_path=temp_db,
    )
    version = build_package_version(database.get_package_by_package_id(package["package_id"], db_path=temp_db), db_path=temp_db)
    return lock_version(version["version_id"], db_path=temp_db)


def test_phase5_schema_upgrade_preserves_phase4_tables(temp_db: Path) -> None:
    with database.get_connection(temp_db) as connection:
        tables = {row["name"] for row in connection.execute("SELECT name FROM sqlite_master WHERE type='table'")}
    assert "package_versions" in tables
    assert "processing_runs" in tables
    assert "document_chunks" in tables
    assert "evidence_records" in tables
    assert "claim_conflicts" in tables


def test_processing_eligibility_blocks_unlocked_failed_missing_and_mutated_versions(temp_db: Path) -> None:
    package = _package(temp_db)
    store_uploaded_files(
        package,
        [UploadCandidate("facts.txt", b"Revenue was $123 million in FY2026.")],
        source_type="other",
        authorization_confirmed=True,
        metadata_by_name={"facts.txt": {"final_category_code": "other"}},
        db_path=temp_db,
    )
    built = build_package_version(package, db_path=temp_db)
    assert not validate_processing_eligibility(built["version_id"], db_path=temp_db).is_eligible
    locked = lock_version(built["version_id"], db_path=temp_db)
    assert validate_processing_eligibility(locked["version_id"], db_path=temp_db).is_eligible
    database.update_package_version(locked["version_id"], {"integrity_status": config.INTEGRITY_FAILED}, db_path=temp_db)
    assert not validate_processing_eligibility(locked["version_id"], db_path=temp_db).is_eligible
    database.update_package_version(locked["version_id"], {"integrity_status": config.INTEGRITY_VERIFIED}, db_path=temp_db)
    doc = database.list_package_version_documents(locked["version_id"], db_path=temp_db)[0]
    root = Path(locked["manifest_path"]).parents[1]
    source = root / doc["relative_package_path"]
    original = source.read_bytes()
    source.write_bytes(b"changed")
    mutated = validate_processing_eligibility(locked["version_id"], db_path=temp_db)
    assert not mutated.is_eligible
    assert any("hash changed" in error.lower() or "size changed" in error.lower() for error in mutated.errors)
    source.write_bytes(original)
    source.unlink()
    missing = validate_processing_eligibility(locked["version_id"], db_path=temp_db)
    assert not missing.is_eligible
    assert any("missing" in error.lower() for error in missing.errors)


def test_pipeline_processes_supported_formats_and_preserves_locators(temp_db: Path) -> None:
    uploads = [
        UploadCandidate("native.pdf", _pdf_bytes(["Page 1 revenue was $100 million in FY2026.", "Page 2 debt was $75 million in FY2026."])),
        UploadCandidate("mixed.pdf", _mixed_pdf_bytes()),
        UploadCandidate("image_only.pdf", _image_only_pdf_bytes()),
        UploadCandidate("notes.docx", _docx_bytes()),
        UploadCandidate("lines.txt", b"Line one\nRevenue was $123 million in FY2026.\nDebt was $40 million."),
        UploadCandidate("table.csv", b"metric,period,value\nrevenue,FY2026,$125 million\nmargin,FY2026,15%\n"),
        UploadCandidate("model.xlsx", _xlsx_bytes()),
        UploadCandidate("macro_model.xlsm", _xlsx_bytes(hidden=False)),
        UploadCandidate("archive.zip", _zip_bytes()),
        UploadCandidate("malformed.pdf", b"%PDF-1.4\nnot actually a valid pdf\n%%EOF"),
    ]
    version = _lock_version_with_uploads(temp_db, uploads)
    original_hashes = {
        doc["document_id"]: sha256_file(Path(version["manifest_path"]).parents[1] / doc["relative_package_path"])
        for doc in database.list_package_version_documents(version["version_id"], db_path=temp_db)
    }
    run = run_processing_pipeline(version["version_id"], db_path=temp_db)
    assert run["status"] == config.PROCESSING_STATUS_COMPLETED_WITH_WARNINGS
    results = database.list_document_processing_results(run["processing_run_id"], db_path=temp_db)
    assert len(results) == len(uploads)
    parser_by_doc = {Path(doc["relative_package_path"]).name: result for doc in database.list_package_version_documents(version["version_id"], db_path=temp_db) for result in results if result["version_document_id"] == doc["document_id"]}
    assert parser_by_doc["native.pdf"]["page_count"] == 2
    assert parser_by_doc["image_only.pdf"]["ocr_required"] == 1
    assert parser_by_doc["archive.zip"]["processing_status"] == config.DOCUMENT_PROCESSING_SKIPPED
    assert parser_by_doc["malformed.pdf"]["processing_status"] == config.DOCUMENT_PROCESSING_FAILED
    pages = database.list_document_pages(run["processing_run_id"], db_path=temp_db)
    assert any(page["page_number"] == 2 and page["extraction_method"] == "NATIVE_TEXT" for page in pages)
    assert any(page["extraction_method"] in {"IMAGE_ONLY", "MIXED"} for page in pages)
    sheets = database.list_document_sheets(run["processing_run_id"], db_path=temp_db)
    assert any(sheet["sheet_name"] == "HiddenData" and sheet["hidden_state"] == "hidden" for sheet in sheets)
    assert any("FORMULA_WITHOUT_CACHED_VALUE" in (sheet["warning_flags"] or "") for sheet in sheets)
    chunks = database.list_document_chunks(run["processing_run_id"], version_id=version["version_id"], db_path=temp_db)
    assert chunks
    assert all(json.loads(chunk["source_locator_json"])["version_document_id"] == chunk["version_document_id"] for chunk in chunks)
    evidence = database.list_evidence_records(run["processing_run_id"], db_path=temp_db)
    assert evidence
    verifications = database.list_citation_verifications(processing_run_id=run["processing_run_id"], db_path=temp_db)
    assert verifications
    assert any(verification["support_status"] == config.VERIFICATION_SUPPORTS for verification in verifications)
    for doc in database.list_package_version_documents(version["version_id"], db_path=temp_db):
        assert sha256_file(Path(version["manifest_path"]).parents[1] / doc["relative_package_path"]) == original_hashes[doc["document_id"]]


def test_retrieval_is_restricted_to_one_version_and_processing_run(temp_db: Path) -> None:
    first = _lock_version_with_uploads(
        temp_db,
        [UploadCandidate("alpha.txt", b"Alpha revenue was $111 million in FY2026."), UploadCandidate("alpha_copy.txt", b"Alpha revenue was $111 million in FY2026. ")],
        ticker="AAA",
    )
    second = _lock_version_with_uploads(
        temp_db,
        [UploadCandidate("beta.txt", b"Beta revenue was $999 million in FY2026.")],
        ticker="BBB",
    )
    first_run = run_processing_pipeline(first["version_id"], db_path=temp_db)
    second_run = run_processing_pipeline(second["version_id"], db_path=temp_db)
    first_results = search_chunks(version_id=first["version_id"], processing_run_id=first_run["processing_run_id"], query="revenue", db_path=temp_db)
    second_results = search_chunks(version_id=second["version_id"], processing_run_id=second_run["processing_run_id"], query="revenue", db_path=temp_db)
    assert first_results
    assert second_results
    assert all(result.version_document_id in {doc["document_id"] for doc in database.list_package_version_documents(first["version_id"], db_path=temp_db)} for result in first_results)
    assert all("999" not in result.chunk_text for result in first_results)
    assert any("999" in result.chunk_text for result in second_results)
    duplicates = database.list_duplicate_groups(first_run["processing_run_id"], db_path=temp_db)
    assert any(group["duplicate_type"] in {"EXACT_CHUNK_DUPLICATE", "NEAR_IDENTICAL_TEXT"} for group in duplicates)
    assert len(first_results) == 1


def test_citation_verification_unsupported_hash_mismatch_and_analyst_evidence(temp_db: Path) -> None:
    version = _lock_version_with_uploads(temp_db, [UploadCandidate("facts.txt", b"Revenue was $123 million in FY2026.")])
    run = run_processing_pipeline(version["version_id"], db_path=temp_db)
    chunk = database.list_document_chunks(run["processing_run_id"], version_id=version["version_id"], db_path=temp_db)[0]
    evidence = database.list_evidence_records(run["processing_run_id"], db_path=temp_db)[0]
    unsupported = dict(evidence)
    unsupported["evidence_id"] = "EVD-UNSUPPORTED"
    unsupported["value"] = 999.0
    unsupported["verification_status"] = config.VERIFICATION_PENDING
    database.create_evidence_record(unsupported, db_path=temp_db)
    verification = verify_evidence_record(unsupported, db_path=temp_db)
    assert verification["support_status"] == config.VERIFICATION_DOES_NOT_SUPPORT
    tampered = dict(evidence)
    tampered["evidence_id"] = "EVD-HASHMISMATCH"
    tampered["source_text"] = "Revenue was $124 million in FY2026."
    tampered["verification_status"] = config.VERIFICATION_PENDING
    database.create_evidence_record(tampered, db_path=temp_db)
    hash_verification = verify_evidence_record(tampered, db_path=temp_db)
    assert hash_verification["support_status"] == config.VERIFICATION_HASH_MISMATCH
    analyst = create_analyst_evidence_from_chunk(
        chunk=chunk,
        evidence_type="OTHER_FACT",
        claim_text="Analyst interpretation tied to the cited chunk.",
        analyst_note="Manual evidence.",
        db_path=temp_db,
    )
    assert analyst["extraction_method"] == "ANALYST_CREATED"
    assert analyst["verification_status"] == config.VERIFICATION_PENDING
    assert analyst["analyst_status"] == config.ANALYST_STATUS_NEEDS_REVIEW
    updated = database.update_evidence_analyst_status(analyst["evidence_id"], config.ANALYST_STATUS_ACCEPTED, "Accepted.", db_path=temp_db)
    assert updated["analyst_status"] == config.ANALYST_STATUS_ACCEPTED


def test_conflict_detection_types_and_no_false_conflict_for_different_metrics(temp_db: Path) -> None:
    version = _lock_version_with_uploads(temp_db, [UploadCandidate("facts.txt", b"Revenue was $123 million in FY2026.")])
    run = run_processing_pipeline(version["version_id"], db_path=temp_db)
    chunk = database.list_document_chunks(run["processing_run_id"], version_id=version["version_id"], db_path=temp_db)[0]
    locator = json.loads(chunk["source_locator_json"])
    base = {
        "processing_run_id": run["processing_run_id"],
        "version_id": version["version_id"],
        "version_document_id": chunk["version_document_id"],
        "normalized_subject": "qxo",
        "source_text": chunk["chunk_text"],
        "page_number": chunk.get("page_number"),
        "sheet_name": chunk.get("sheet_name"),
        "cell_or_row_range": chunk.get("row_range"),
        "section_heading": chunk.get("section_heading"),
        "extraction_method": "TEST",
        "confidence": "High",
        "verification_status": config.VERIFICATION_PENDING,
        "analyst_status": config.ANALYST_STATUS_UNREVIEWED,
        "analyst_note": "",
        "source_locator_json": json.dumps({**locator, "chunk_id": chunk["chunk_id"]}, sort_keys=True),
        "source_text_hash": None,
        "created_by": "test",
        "created_at": database.utc_now_iso(),
        "updated_at": database.utc_now_iso(),
    }
    for evidence_id, metric, value, unit, claim in [
        ("EVD-C1", "eps", 1.25, None, "EPS estimate was 1.25 in FY2026."),
        ("EVD-C2", "eps", 1.55, None, "EPS estimate was 1.55 in FY2026."),
        ("EVD-C3", "debt", 1.55, None, "Debt was 1.55 in FY2026."),
        ("EVD-C4", "margin", 15.0, "%", "GAAP margin was 15% in FY2026."),
        ("EVD-C5", "margin", 15.0, "%", "Adjusted margin was 15% in FY2026."),
    ]:
        database.create_evidence_record(
            {
                **base,
                "evidence_id": evidence_id,
                "evidence_type": "ANALYST_ESTIMATE" if metric == "eps" else "OTHER_FACT",
                "claim_text": claim,
                "metric_name": metric,
                "value": value,
                "unit": unit,
                "currency": None,
                "period": "FY2026",
                "scenario": None,
                "direction": None,
            },
            db_path=temp_db,
        )
    conflicts = detect_claim_conflicts(processing_run_id=run["processing_run_id"], db_path=temp_db)
    conflict_types = {conflict["conflict_type"] for conflict in conflicts}
    assert "FORECAST_DISAGREEMENT" in conflict_types
    assert "GAAP_ADJUSTED_MISMATCH" in conflict_types
    assert not any(conflict["metric"] == "debt" and conflict["conflict_type"] == "VALUE_DIFFERENCE" for conflict in conflicts)


def test_pdf_parser_edge_cases_without_ocr(temp_db: Path, tmp_path: Path) -> None:
    version_doc = {
        "document_id": "VDOC-PDF",
        "original_document_id": "DOC-PDF",
        "title": "Edge PDF",
        "relative_package_path": "edge.pdf",
    }
    native_path = tmp_path / "native.pdf"
    native_path.write_bytes(_pdf_bytes(["Native revenue was $10 million.", "Second page margin was 12%."]))
    parsed = parse_version_document(version_doc=version_doc, source_path=native_path, version_id="VER", processing_run_id="RUN", ocr_enabled=False)
    assert parsed.status == config.DOCUMENT_PROCESSING_SUCCESS
    assert parsed.page_count == 2
    image_path = tmp_path / "image.pdf"
    image_path.write_bytes(_image_only_pdf_bytes())
    image_parsed = parse_version_document(version_doc=version_doc, source_path=image_path, version_id="VER", processing_run_id="RUN2", ocr_enabled=False)
    assert image_parsed.ocr_required
    assert image_parsed.status == config.DOCUMENT_PROCESSING_PARTIAL
    empty_path = tmp_path / "empty.pdf"
    import fitz

    empty_path.write_bytes(
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[]/Count 0>>endobj\n"
        b"trailer<</Root 1 0 R>>\n%%EOF\n"
    )
    empty_parsed = parse_version_document(version_doc=version_doc, source_path=empty_path, version_id="VER", processing_run_id="RUN3", ocr_enabled=False)
    assert empty_parsed.status == config.DOCUMENT_PROCESSING_PARTIAL
    malformed_path = tmp_path / "bad.pdf"
    malformed_path.write_bytes(b"%PDF-1.4\nbad")
    malformed = parse_version_document(version_doc=version_doc, source_path=malformed_path, version_id="VER", processing_run_id="RUN4", ocr_enabled=False)
    assert malformed.status == config.DOCUMENT_PROCESSING_FAILED


def test_phase5_streamlit_empty_states_load() -> None:
    for path in [
        "app/Home.py",
        "app/pages/1_New_Research_Package.py",
        "app/pages/2_Document_Collection.py",
        "app/pages/3_Package_Review.py",
        "app/pages/4_Investment_Analysis.py",
        "app/pages/5_Generated_Reports.py",
    ]:
        app = AppTest.from_file(path, default_timeout=10)
        app.run()
        assert not list(app.exception)
